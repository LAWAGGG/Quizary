import re
import io
import os
import uuid
from collections import Counter

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status
from sqlalchemy.orm import Session

from app.database import get_db
from app.dependencies import verify_form_owner
from app.models.form import Form
from app.models.image import Image
from app.models.question import Question, QuestionType
from app.models.question_option import QuestionOption
from app.services.points import distribute_quiz_points
from app.utils import UPLOAD_DIR, now_wib

router = APIRouter(tags=["import"])

WORD_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
A_NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"

ANSWER_RE = re.compile(r'(?:Kunci\s*)?(?:Jawaban|Jawab|Answer)\s*[:\-]?\s*([A-Da-d])\b', re.IGNORECASE)
NUMBERED_RE = re.compile(r'\d+[\.\)]\s*(.+)')
OPTION_INLINE_RE = re.compile(r'(?:^|\s)([A-Da-d])[\.\)]\s*(.*?)(?=\s+[A-Da-d][\.\)]|$)')


# ============================================================
# Path 1: plain text paste ("POST /import/text") — tidak berubah,
# tetap dipakai untuk teks yang diketik manual (tidak ada info numbering Word)
# ============================================================
def _parse_text(raw: str) -> list[dict]:
    questions: list[dict] = []
    blocks = re.split(r'\n\s*(?=\d+[\.\)])', raw.strip())
    for block in blocks:
        block = block.strip()
        if not block:
            continue

        q_text = None
        options: list[dict] = []
        answer_letter = None

        for line in block.split("\n"):
            line = line.strip()
            if not line:
                continue
            m = NUMBERED_RE.match(line)
            if m:
                q_text = m.group(1).strip()
                continue
            m = ANSWER_RE.match(line)
            if m:
                answer_letter = m.group(1).upper()
                continue
            for om in OPTION_INLINE_RE.finditer(line):
                text = om.group(2).strip()
                if text:
                    options.append({"letter": om.group(1).upper(), "text": text})

        if q_text:
            questions.append({
                "question_text": q_text,
                "options": [
                    {"text": opt["text"], "is_correct": opt["letter"] == answer_letter}
                    for opt in options
                ],
            })

    return questions


# ============================================================
# Path 2: import .docx — DIPERBAIKI supaya bisa baca native Word
# numbered list (numPr), bukan cuma angka yang diketik manual
# ============================================================
def _para_images(p) -> list[tuple[str, bytes]]:
    """Ambil semua gambar yang tertanam di dalam sebuah paragraf docx.

    Docx adalah file ZIP; gambar dirender lewat relasi rId pada elemen
    ``<a:blip r:embed="rIdN">``. Blob biner diambil langsung dari part —
    TANPA base64, sehingga ukuran file asli tidak membengkak 33%.
    """
    imgs: list[tuple[str, bytes]] = []
    for blip in p._p.findall(f".//{A_NS}blip"):
        rid = blip.get(f"{R_NS}embed")
        if not rid:
            continue
        rel = p.part.rels.get(rid)
        if rel is None or rel.is_external:
            continue
        part = rel.target_part
        ext = os.path.splitext(str(part.partname))[1].lower()
        if not ext:
            ext = ".png"
        imgs.append((ext, part.blob))
    return imgs


def _extract_docx_items(doc) -> list[tuple[str, str | None, list]]:
    """
    Kembalikan list (text, num_id, images) per paragraf yang punya makna.
    num_id None berarti paragraf itu TIDAK memakai numbering otomatis Word
    (baik karena memang teks biasa, maupun karena angkanya diketik manual —
    dua kasus ini tetap bisa dibedakan lewat regex NUMBERED_RE di caller).

    Paragraf text kosong TIDAK di-skip bila ia membawa numbering (nomor soal
    auto Word) atau gambar (stem soal berupa gambar) — keduanya krusial.
    """
    items = []
    for p in doc.paragraphs:
        text = p.text.strip()
        numPr = p._p.find(f".//{WORD_NS}numPr")
        num_id = None
        if numPr is not None:
            numId_el = numPr.find(f"{WORD_NS}numId")
            if numId_el is not None:
                num_id = numId_el.get(f"{WORD_NS}val")
        imgs = _para_images(p)
        if not text and num_id is None and not imgs:
            continue
        items.append((text, num_id, imgs))
    return items


def _parse_docx_items(items: list[tuple[str, str | None, list]]) -> list[dict]:
    if not items:
        return []

    # Cari numId yang paling sering muncul -> itu kemungkinan besar
    # list penomoran SOAL (karena berulang 1x per soal di sepanjang dokumen),
    # numId lain yang muncul lokal biasanya sub-list opsi jawaban.
    numid_counts = Counter(nid for _, nid, _ in items if nid is not None)
    question_num_id = numid_counts.most_common(1)[0][0] if numid_counts else None

    questions: list[dict] = []
    current: dict | None = None
    active_option_num_id: str | None = None
    option_letters = iter("ABCDEFGHIJ")

    def flush():
        nonlocal current
        if current and (current["question_text"] or current["images"]):
            questions.append(current)
        current = None

    def attach_imgs(target: dict, imgs: list):
        for ext, blob in imgs:
            target["images"].append({"ext": ext, "blob": blob})

    def start_question(text: str, imgs: list):
        nonlocal current, active_option_num_id, option_letters
        flush()
        current = {"question_text": text.strip(), "options": [], "answer_letter": None, "images": []}
        active_option_num_id = None
        option_letters = iter("ABCDEFGHIJ")
        attach_imgs(current, imgs)

    for text, num_id, imgs in items:
        # 1) angka diketik manual "1. ..." / "1) ..." (bisa juga muncul di dokumen
        #    yang campur: sebagian numbering manual, sebagian native Word list)
        m = NUMBERED_RE.match(text)
        if m and (num_id is None or num_id == question_num_id):
            start_question(m.group(1), imgs)
            continue

        # 2) baris "Jawaban: X"
        m = ANSWER_RE.match(text)
        if m and current is not None:
            current["answer_letter"] = m.group(1).upper()
            continue

        # 3) opsi manual "A. ..." (termasuk multi-kolom satu baris "A. x D. y")
        found_option = False
        if current is not None:
            for om in OPTION_INLINE_RE.finditer(text):
                opt_text = om.group(2).strip()
                if opt_text:
                    opt = {"letter": om.group(1).upper(), "text": opt_text, "images": []}
                    attach_imgs(opt, imgs)
                    current["options"].append(opt)
                    found_option = True
        if found_option:
            continue

        # 4) native Word numbered list, numId = numId soal (mayoritas) -> soal baru.
        #    Paragraf bisa text kosong (nomor auto Word ada di paragraf terpisah).
        if num_id is not None and num_id == question_num_id:
            start_question(text, imgs)
            continue

        # 5) native Word numbered list, numId BEDA -> anggap sub-list opsi jawaban
        if num_id is not None and current is not None:
            if active_option_num_id != num_id:
                active_option_num_id = num_id
                option_letters = iter("ABCDEFGHIJ")
            letter = next(option_letters, "?")
            opt = {"letter": letter, "text": text, "images": []}
            attach_imgs(opt, imgs)
            current["options"].append(opt)
            continue

        # 6) tidak ada numbering & tidak match pola apapun -> baris lanjutan
        #    (soal/opsi yang wrap ke baris baru, atau stem soal berupa gambar),
        #    gabungkan ke item terakhir
        if current is not None:
            if current["options"]:
                last_opt = current["options"][-1]
                if text:
                    last_opt["text"] += " " + text
                attach_imgs(last_opt, imgs)
            else:
                if text:
                    current["question_text"] += " " + text
                attach_imgs(current, imgs)
            continue

        # 7) belum ada soal terbuka (misal judul dokumen di baris pertama) -> lewati

    flush()

    # finalisasi is_correct
    result = []
    for q in questions:
        answer_letter = q.pop("answer_letter", None)
        result.append({
            "question_text": q["question_text"],
            "images": q["images"],
            "options": [
                {"text": o["text"], "is_correct": o["letter"] == answer_letter, "images": o["images"]}
                for o in q["options"]
            ],
        })
    return result


_WEB_IMG_EXT = {".png", ".jpg", ".jpeg", ".gif", ".webp"}


def _save_blob(ext: str, blob: bytes) -> str:
    """Tulis blob biner docx ke disk uploads, kembalikan path relatif.

    Hanya format web (dapat dirender <img>) yang disimpan; format vektor
    docx seperti .emf/.wmf di-skip — tidak bisa dirender browser.
    """
    if ext not in _WEB_IMG_EXT:
        return ""
    filename = f"{uuid.uuid4().hex}{ext}"
    dest = os.path.join(UPLOAD_DIR, "question-images", filename)
    os.makedirs(os.path.dirname(dest), exist_ok=True)
    with open(dest, "wb") as f:
        f.write(blob)
    return f"question-images/{filename}"


@router.post("/forms/{form_id}/import/docx", status_code=201)
async def import_docx(
    form: Form = Depends(verify_form_owner),
    db: Session = Depends(get_db),
    file: UploadFile = File(...),
):
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Only .docx files are supported",
        )
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_501_NOT_IMPLEMENTED,
            detail="python-docx is not installed",
        )

    content = await file.read()
    doc = Document(io.BytesIO(content))

    items = _extract_docx_items(doc)
    parsed = _parse_docx_items(items)

    if not parsed:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="No questions could be imported, check document format",
        )

    max_order = (
        db.query(Question.order_index)
        .filter(Question.form_id == form.id)
        .order_by(Question.order_index.desc())
        .first()
    )
    next_order = (max_order[0] + 1) if max_order else 0
    now = now_wib()
    count = 0

    for q_data in parsed:
        has_options = bool(q_data["options"])
        correct_count = sum(1 for o in q_data["options"] if o["is_correct"])

        if has_options and correct_count > 1:
            q_type = QuestionType.checkbox
        elif has_options:
            q_type = QuestionType.multiple_choice
        else:
            q_type = QuestionType.essay

        q = Question(
            form_id=form.id,
            type=q_type,
            question_text=q_data["question_text"],
            points=0 if form.type.value == "quiz" else 1,
            order_index=next_order,
            created_at=now,
        )
        db.add(q)
        db.flush()

        for i, img in enumerate(q_data["images"]):
            path = _save_blob(img["ext"], img["blob"])
            if path:
                db.add(Image(
                    question_id=q.id,
                    path=path,
                    order_index=i,
                    created_at=now,
                ))

        opt_rows = []
        for i, opt in enumerate(q_data["options"]):
            opt_rows.append(QuestionOption(
                question_id=q.id,
                option_text=opt["text"],
                is_correct=opt["is_correct"],
                order_index=i,
            ))
            db.add(opt_rows[-1])
        db.flush()

        for opt_row, opt in zip(opt_rows, q_data["options"]):
            for j, img in enumerate(opt["images"]):
                path = _save_blob(img["ext"], img["blob"])
                if path:
                    db.add(Image(
                        option_id=opt_row.id,
                        path=path,
                        order_index=j,
                        created_at=now,
                    ))

        next_order += 1
        count += 1

    distribute_quiz_points(form.id, db)
    db.commit()
    return {"message": f"{count} question(s) imported successfully", "imported_count": count}


if __name__ == "__main__":
    import os
    from docx import Document  # type: ignore

    # fixture: plain text (manual numbering, essay + MCQ multi-column)
    s = """1. Tes essay saja apakah terbaca
2. Soal pilihan ganda tanpa kunci jawaban
A. 252 cara\t\tD. 258 cara
C. 256 cara
3. Opsi gambar/placeholder
A.
D.
"""
    qs = _parse_text(s)
    assert len(qs) == 3, qs
    assert qs[0]["options"] == [], "essay harus disimpan tanpa opsi"
    assert [o["text"] for o in qs[1]["options"]] == ["252 cara", "258 cara", "256 cara"]
    assert qs[2]["options"] == [], "opsi kosong harus dibuang"

    # fixture: docx dengan native Word numbering (numPr)
    sample = os.path.join(os.path.dirname(__file__), "../../../tmp_test/soal mtk.docx")
    if os.path.exists(sample):
        doc = Document(sample)
        parsed = _parse_docx_items(_extract_docx_items(doc))
        mc = [q for q in parsed if q["options"]]
        assert len(mc) >= 2, mc
        assert all(len(q["options"]) == 5 for q in mc), [len(q["options"]) for q in mc]
        assert any(len(q["options"]) == 0 for q in parsed), "harus ada essay (soal gambar/no.3)"
        print(f"ok docx; {len(parsed)} soal, {len(mc)} MCQ (masing-masing 5 opsi)")

    # fixture nyata: stem soal berupa gambar, nomor soal auto-numbering di
    # paragraf kosong, opsi teks tanpa numbering -> 20 soal MCQ, 20 gambar stem
    real = os.path.join(os.path.dirname(__file__), "../../../ULANGAN HARIAN BAHASA INDONESIA.docx")
    if os.path.exists(real):
        doc = Document(real)
        parsed = _parse_docx_items(_extract_docx_items(doc))
        assert len(parsed) == 20, f"harus 20 soal, dapat {len(parsed)}"
        assert all(len(q["images"]) == 1 for q in parsed), "tiap soal harus punya 1 gambar stem"
        assert all(q["images"][0]["ext"] == ".png" for q in parsed), "stem harus PNG"
        assert all(3 <= len(q["options"]) <= 5 for q in parsed), "opsi 3-5 per soal"
        assert all(o["text"] for q in parsed for o in q["options"]), "ada opsi kosong"
        print(f"ok real docx; {len(parsed)} soal, 20 gambar stem PNG, opsi 3-5 per soal")
    print("done")